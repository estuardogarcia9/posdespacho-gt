"""
Parsear Posdespacho AMM
=======================
Lee un ZIP diario del AMM (ej: PD20260408.zip), extrae los documentos Word
embebidos en las hojas RSO1 y RSO2, parsea los eventos y los agrega a un
Excel acumulativo (posdespachos_acumulado.xlsx).

Uso:
    python parsear_posdespacho.py PD20260408.zip
    python parsear_posdespacho.py               (busca el ZIP más reciente en la carpeta)
"""

import sys
import re
import zipfile
import io
import os
from datetime import datetime, date, time, timedelta
from pathlib import Path

from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ─── Configuración ────────────────────────────────────────────────────────────
# Rutas independientes del usuario. Se pueden sobreescribir con variables de
# entorno (POSDESPACHO_ZIPS / POSDESPACHO_EXCEL) para casos especiales.

BASE            = Path(__file__).parent
CARPETA_ZIPS    = Path(os.environ.get("POSDESPACHO_ZIPS", Path.home() / "Downloads"))  # donde caen los ZIPs del AMM
EXCEL_SALIDA    = Path(os.environ.get("POSDESPACHO_EXCEL", BASE / "posdespachos_acumulado.xlsx"))

# ─── Columnas del Excel de salida ─────────────────────────────────────────────

COLUMNAS = [
    "Estado del evento",        # A  1
    "¿Requiere Investigar?",    # B  2
    "ID",                       # C  3
    "ID Único",                 # D  4
    "Categoría",                # E  5
    "Narrativa del AMM",        # F  6
    "Fecha inicial",            # G  7
    "Fecha final",              # H  8
    "Hora Inicial",             # I  9
    "Hora Final",               # J  10
    "Duración del evento en horas",           # K  11
    "Duración del evento en horas y minutos", # L  12
    "Detonante",                # M  13
    "Sub-detonante",            # N  14
    "Tipo de afectación",       # O  15
    "Departamento",             # P  16
    "Propietario",              # Q  17
    "Tipo de Propietario",      # R  18
    "Activo indicado",          # S  19
    "Activo",                   # T  20
    "Clasificación",            # U  21
    "Dirección",                # V  22
    "Especificación",           # W  23
    "Carga afectada",           # X  24
    "Asociación de Eventos",    # Y  25
    "Código de Mantenimiento",  # Z  26
    "Colateralidad",            # AA 27
    "Comentarios",              # AB 28
    "Imagen",                   # AC 29
    "Duración de eventos abiertos", # AD 30
    "Narrativa",                # AE 31
    "Estado del activo",        # AF 32
    "Len()",                    # AG 33
    "",                         # AH 34
    "Fuente",                   # AI 35
]

# ─── Helpers de parsing ───────────────────────────────────────────────────────

RE_HORAS       = re.compile(r'de\s+(\d{1,2}:\d{2})\s+a\s+(\d{1,2}:\d{2})\s*(?:horas?|,|\.)', re.IGNORECASE)
RE_HORA_UNICA  = re.compile(r'a\s+las?\s*(\d{1,2}:\d{2})\s+horas', re.IGNORECASE)
RE_CODIGO_MANT = re.compile(r'(AMM-[A-Z]{3}\d{2}-[A-Z]{3}\d+)', re.IGNORECASE)
RE_VOLTAJE     = re.compile(r'(\d{2,3})\s*kV', re.IGNORECASE)


def parse_hora(h_str: str) -> time | None:
    """Convierte 'HH:MM' o 'H:MM' a time object. '24:00' -> time(0,0)."""
    try:
        parts = h_str.strip().split(":")
        hora = int(parts[0])
        minuto = int(parts[1])
        if hora == 24:
            hora = 0
        return time(hora, minuto)
    except Exception:
        return None


def calcular_duracion(h_ini: time, h_fin: time):
    """Retorna (horas_float, texto_duracion) o (None, None)."""
    if h_ini is None or h_fin is None:
        return None, None
    dt_ini = datetime.combine(date.today(), h_ini)
    dt_fin = datetime.combine(date.today(), h_fin)
    if dt_fin < dt_ini:          # cruza medianoche
        dt_fin += timedelta(days=1)
    delta = dt_fin - dt_ini
    total_min = int(delta.total_seconds() // 60)
    dias  = total_min // (24 * 60)
    horas = (total_min % (24 * 60)) // 60
    mins  = total_min % 60
    horas_float = delta.total_seconds() / 3600
    texto = f"{dias} días, {horas} horas y {mins} minutos"
    return horas_float, texto


def inferir_detonante(texto: str, categoria: str):
    """
    Devuelve (detonante, sub_detonante) basado en palabras clave del texto.
    """
    t = texto.lower()

    # Interconexiones siempre son Operación estable
    if categoria == "Interconexiones":
        return "Operación\nestable", "Interconexión\ncerrada"

    # Indisponibilidad forzada (verificar antes de programada)
    forzada_kws = ["disparo", "disparó", "dispararon", "abierta por disparo",
                   "abierto por disparo", "forzada", "falla", "trabajos forzados",
                   "reparación", "reparacion"]
    if any(kw in t for kw in forzada_kws):
        # Excepción: si también tiene código de mantenimiento AMM-, es programada
        if RE_CODIGO_MANT.search(texto):
            return "Indisponibilidad\nprogramada", "Mantenimiento\nprogramado"
        return "Indisponibilidad\nforzada", "Disparo"

    # Indisponibilidad programada
    mant_code = RE_CODIGO_MANT.search(texto)
    progr_kws = ["mantenimiento", "se abri"]
    # "programado/a" solo si va precedido por algo positivo (no "no programado")
    tiene_programado = bool(re.search(r'(?<!no\s)\bprogramad[ao]\b', t))
    tiene_desenergiz = "desenergiz" in t
    if mant_code or tiene_programado or tiene_desenergiz or any(kw in t for kw in progr_kws):
        return "Indisponibilidad\nprogramada", "Mantenimiento\nprogramado"

    # Maniobras
    if any(kw in t for kw in ["maniobra", "traslad", "conecta", "desconecta",
                               "se traslad", "abiertos interruptores"]):
        return "Maniobras", "Maniobra"

    return None, None


def inferir_clasificacion(texto: str) -> str | None:
    """Infiere la clasificación (LT 69 kV, INT 230 kV, etc.) del texto."""
    t_lower = texto.lower()
    voltaje_match = RE_VOLTAJE.search(texto)
    kv = voltaje_match.group(1) if voltaje_match else None

    if not kv:
        return None

    if "interconex" in t_lower:
        return f"INT {kv} kV"
    elif any(w in t_lower for w in ["línea", "linea"]):
        return f"LT {kv} kV"
    elif any(w in t_lower for w in ["transformador", "trafo"]):
        return f"TX {kv} kV"
    elif any(w in t_lower for w in ["subestación", "subestacion", "barra"]):
        return f"SE {kv} kV"
    else:
        return f"LT {kv} kV"


def extraer_activo_indicado(texto: str) -> str | None:
    """
    Extrae el nombre del activo de la narrativa.
    Ej: "la línea 69 kV Escuintla – El Jocote" → "LT ESCUINTLA-EL JOCOTE"
    """
    t = texto

    # Interconexión
    m = re.search(r'interconexi[oó]n\s+\d+\s*kV\s+Guatemala[^–-]*[–-]\s*([A-Za-záéíóúñüÁÉÍÓÚÑÜ\s]+)', t, re.IGNORECASE)
    if m:
        pais = m.group(1).strip().split()[0].upper()
        return f"INT GUATEMALA-{pais}"

    # Línea kV A – B o A - B
    m = re.search(
        r'l[ií]nea\s+[\d.]+\s*kV\s+([A-Za-záéíóúñüÁÉÍÓÚÑÜ\s\.]+?)\s*[–\-]\s*([A-Za-záéíóúñüÁÉÍÓÚÑÜ\s\.]+?)(?:\s+de\s+\d|\s+a\s+\d|,|\.|$)',
        t, re.IGNORECASE
    )
    if m:
        a = re.sub(r'\s+', ' ', m.group(1)).strip().upper()
        b = re.sub(r'\s+', ' ', m.group(2)).strip().upper()
        return f"LT {a}-{b}"

    return None


# ─── Extracción de documentos Word desde el ZIP del AMM ───────────────────────

def extraer_docx_desde_zip(zip_path: Path) -> tuple[bytes, bytes]:
    """
    Abre el ZIP del AMM → abre el Excel interno → extrae los dos DOCX embebidos.
    Retorna (rso1_docx_bytes, rso2_docx_bytes).
    """
    with zipfile.ZipFile(zip_path, 'r') as z_outer:
        # Encontrar el archivo xlsx dentro del ZIP
        xlsx_names = [n for n in z_outer.namelist() if n.lower().endswith('.xlsx')]
        if not xlsx_names:
            raise ValueError(f"No se encontró .xlsx dentro de {zip_path.name}")
        xlsx_bytes = z_outer.read(xlsx_names[0])

    with zipfile.ZipFile(io.BytesIO(xlsx_bytes), 'r') as z_xlsx:
        # Los embeds son: xl/embeddings/Microsoft_Word_Document.docx  (RSO1)
        #                 xl/embeddings/Microsoft_Word_Document1.docx (RSO2)
        embed_names = [n for n in z_xlsx.namelist()
                       if 'embeddings' in n and n.lower().endswith('.docx')]
        embed_names.sort()   # Document.docx < Document1.docx
        if len(embed_names) < 2:
            raise ValueError(f"Se esperaban 2 DOCX embebidos, se encontraron {len(embed_names)}")
        rso1_bytes = z_xlsx.read(embed_names[0])
        rso2_bytes = z_xlsx.read(embed_names[1])

    return rso1_bytes, rso2_bytes


# ─── Parseo de cada sección de los DOCX ──────────────────────────────────────

def parsear_rso1(docx_bytes: bytes, fecha: date) -> list[dict]:
    """Extrae eventos de RSO1 (Aspectos operativos + Interconexiones)."""
    doc = Document(io.BytesIO(docx_bytes))
    eventos = []

    # Identificar secciones por títulos
    seccion_actual = None

    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        # Detectar cambio de sección
        texto_upper = texto.upper()
        if "ASPECTOS OPERATIVOS" in texto_upper:
            seccion_actual = "Aspectos\noperativos"
            continue
        if "SISTEMA ELECTRICO REGIONAL" in texto_upper or "SISTEMA ELÉCTRICO REGIONAL" in texto_upper:
            seccion_actual = "Interconexiones"
            continue
        if "DATOS HIDROLOGICOS" in texto_upper or "DATOS HIDROLÓGICOS" in texto_upper:
            seccion_actual = None   # tabla, no párrafos de eventos
            continue

        # Solo procesar párrafos de evento
        if seccion_actual is None:
            continue
        if p.style.name in ('Normal',) and texto_upper == texto_upper:
            # posible título menor, skip si es todo mayúsculas y corto
            if len(texto) < 60 and texto == texto.upper():
                continue

        # Crear evento
        evento = _crear_evento_base(texto, fecha, seccion_actual, "RSO1")
        eventos.append(evento)

    return eventos


def parsear_rso2(docx_bytes: bytes, fecha: date) -> list[dict]:
    """Extrae eventos de RSO2 (Sistemas de transmisión)."""
    doc = Document(io.BytesIO(docx_bytes))
    eventos = []

    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        texto_upper = texto.upper()
        # Saltar títulos de sección
        if "INFORMACIÓN DEL SISTEMA" in texto_upper or "INFORMACION DEL SISTEMA" in texto_upper:
            continue
        if texto_upper == texto_upper and len(texto) < 60 and re.match(r'^[A-ZÁÉÍÓÚÑÜ\s]+$', texto):
            continue

        evento = _crear_evento_base(texto, fecha, "Sistemas de\ntransmisión", "RSO2")
        eventos.append(evento)

    return eventos


def _crear_evento_base(texto: str, fecha: date, categoria: str, fuente: str) -> dict:
    """Crea un diccionario de evento con todos los campos posibles."""
    e = {col: None for col in COLUMNAS}

    e["Estado del evento"] = "Abierto"
    e["Categoría"]         = categoria
    e["Narrativa del AMM"] = "•  " + texto
    e["Fecha inicial"]     = fecha
    e["Estado del activo"] = "Perturbación"
    e["Fuente"]            = fuente
    e["Len()"]             = len(texto)

    # Extraer horas (primera ocurrencia del patrón "de HH:MM a HH:MM horas")
    m_horas = RE_HORAS.search(texto)
    if m_horas:
        h_ini = parse_hora(m_horas.group(1))
        h_fin = parse_hora(m_horas.group(2))
        e["Hora Inicial"] = h_ini
        e["Hora Final"]   = h_fin
        dur_h, dur_txt    = calcular_duracion(h_ini, h_fin)
        e["Duración del evento en horas"]             = dur_h
        e["Duración del evento en horas y minutos"]   = dur_txt
        e["Fecha final"] = fecha

    # Código de mantenimiento
    m_cod = RE_CODIGO_MANT.search(texto)
    if m_cod:
        e["Código de Mantenimiento"] = m_cod.group(1).upper()

    # Detonante
    det, sub = inferir_detonante(texto, categoria)
    e["Detonante"]    = det
    e["Sub-detonante"] = sub

    # Clasificación
    e["Clasificación"] = inferir_clasificacion(texto)

    # Activo indicado
    e["Activo indicado"] = extraer_activo_indicado(texto)

    return e


# ─── Escritura en Excel ───────────────────────────────────────────────────────

COLOR_HEADER = "1F3864"   # azul oscuro
COLOR_RSO1   = "E2EFDA"   # verde claro
COLOR_RSO2   = "DEEAF1"   # azul claro
COLOR_ALT    = "F2F2F2"   # gris muy suave

def _header_style(cell):
    cell.font      = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    cell.fill      = PatternFill("solid", fgColor=COLOR_HEADER)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def _borde_fino():
    lado = Side(style="thin", color="CCCCCC")
    return Border(left=lado, right=lado, top=lado, bottom=lado)


def guardar_excel(eventos: list[dict], ruta_salida: Path, proximo_id: int):
    """
    Escribe (o agrega) eventos al Excel acumulativo.
    Si el archivo no existe, lo crea con encabezados y formato.
    """
    if ruta_salida.exists():
        wb = openpyxl.load_workbook(ruta_salida)
        ws = wb.active
        # Determinar el último ID utilizado
        ultima_fila = ws.max_row
        # Buscar el último ID en col C (col 3)
        for fila in range(ultima_fila, 1, -1):
            val = ws.cell(row=fila, column=3).value
            if val is not None and str(val).isdigit():
                proximo_id = int(val) + 1
                break
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Posdespachos"

        # Encabezados
        for col_idx, nombre in enumerate(COLUMNAS, start=1):
            cell = ws.cell(row=1, column=col_idx, value=nombre)
            _header_style(cell)

        # Anchos de columna
        anchos = {
            1: 12, 2: 12, 3: 8, 4: 10, 5: 16, 6: 70, 7: 13, 8: 13,
            9: 10, 10: 10, 11: 12, 12: 22, 13: 18, 14: 20, 15: 16,
            16: 20, 17: 40, 18: 14, 19: 35, 20: 35, 21: 14, 22: 10,
            26: 22, 31: 50, 32: 14, 33: 8, 35: 8
        }
        for col_idx, ancho in anchos.items():
            ws.column_dimensions[get_column_letter(col_idx)].width = ancho

        ws.freeze_panes = "A2"
        ws.row_dimensions[1].height = 30
        ultima_fila = 1

    # Agregar filas
    for i, ev in enumerate(eventos):
        fila = ultima_fila + 1 + i
        ev["ID"]       = proximo_id
        ev["ID Único"] = f"{proximo_id}.1"
        proximo_id    += 1

        for col_idx, col_name in enumerate(COLUMNAS, start=1):
            val = ev.get(col_name)
            cell = ws.cell(row=fila, column=col_idx, value=val)

            # Formato de fecha y hora
            if col_name in ("Fecha inicial", "Fecha final") and isinstance(val, date):
                cell.number_format = "DD/MM/YYYY"
            elif col_name in ("Hora Inicial", "Hora Final") and isinstance(val, time):
                cell.number_format = "HH:MM"
            elif col_name == "Duración del evento en horas" and val is not None:
                cell.number_format = "0.00"

            # Color de fila según fuente
            fuente = ev.get("Fuente", "")
            if fuente == "RSO1":
                fill_color = COLOR_RSO1
            elif fuente == "RSO2":
                fill_color = COLOR_RSO2
            else:
                fill_color = COLOR_ALT if fila % 2 == 0 else "FFFFFF"

            cell.fill      = PatternFill("solid", fgColor=fill_color)
            cell.border    = _borde_fino()
            cell.alignment = Alignment(vertical="top", wrap_text=(col_idx == 6))

    try:
        wb.save(ruta_salida)
    except PermissionError:
        alt = ruta_salida.with_stem(ruta_salida.stem + "_PENDIENTE")
        wb.save(alt)
        print(f"  AVISO: El archivo principal esta abierto en Excel.")
        print(f"  Se guardo una copia en: {alt.name}")
        print(f"  Cierra Excel y renombra/copia ese archivo a: {ruta_salida.name}")
    return proximo_id - 1  # último ID utilizado


# ─── Obtener la fecha desde el nombre del ZIP ─────────────────────────────────

def fecha_desde_zip(zip_path: Path) -> date:
    """PD20260408.zip → date(2026, 4, 8)"""
    m = re.search(r'PD(\d{4})(\d{2})(\d{2})', zip_path.name, re.IGNORECASE)
    if m:
        return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    # Fallback: hoy
    return date.today()


# ─── Punto de entrada ─────────────────────────────────────────────────────────

def main():
    # Determinar qué ZIP procesar
    if len(sys.argv) > 1:
        zip_path = Path(sys.argv[1])
        if not zip_path.is_absolute():
            zip_path = CARPETA_ZIPS / zip_path
    else:
        # Buscar el ZIP más reciente en la carpeta
        zips = sorted(CARPETA_ZIPS.glob("PD????????.zip"))
        if not zips:
            print(f"No se encontró ningún ZIP en {CARPETA_ZIPS}")
            sys.exit(1)
        zip_path = zips[-1]

    if not zip_path.exists():
        print(f"Archivo no encontrado: {zip_path}")
        sys.exit(1)

    print(f"Procesando: {zip_path.name}")

    # Fecha del posdespacho
    fecha = fecha_desde_zip(zip_path)
    print(f"Fecha del posdespacho: {fecha.strftime('%d/%m/%Y')}")

    # Extraer los DOCX
    print("Extrayendo documentos Word embebidos...")
    rso1_bytes, rso2_bytes = extraer_docx_desde_zip(zip_path)

    # Parsear eventos
    print("Parseando RSO1 (Aspectos operativos + Interconexiones)...")
    eventos_rso1 = parsear_rso1(rso1_bytes, fecha)
    print(f"  -> {len(eventos_rso1)} eventos")

    print("Parseando RSO2 (Sistemas de transmision)...")
    eventos_rso2 = parsear_rso2(rso2_bytes, fecha)
    print(f"  -> {len(eventos_rso2)} eventos")

    todos = eventos_rso1 + eventos_rso2

    # Determinar próximo ID
    proximo_id = 1
    if EXCEL_SALIDA.exists():
        wb_tmp = openpyxl.load_workbook(EXCEL_SALIDA, read_only=True)
        ws_tmp = wb_tmp.active
        for fila in range(ws_tmp.max_row, 1, -1):
            val = ws_tmp.cell(row=fila, column=3).value
            if val is not None and str(val).isdigit():
                proximo_id = int(val) + 1
                break
        wb_tmp.close()

    # Guardar
    EXCEL_SALIDA.parent.mkdir(parents=True, exist_ok=True)
    print(f"Guardando en: {EXCEL_SALIDA}")
    ultimo_id = guardar_excel(todos, EXCEL_SALIDA, proximo_id)

    print(f"\nOK: {len(todos)} eventos agregados (IDs {proximo_id} - {ultimo_id})")
    print(f"OK: Excel guardado en: {EXCEL_SALIDA}")


if __name__ == "__main__":
    main()
